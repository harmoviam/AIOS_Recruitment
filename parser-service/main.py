"""Resume parser microservice.

Text extraction: pdfplumber (PDF), python-docx (DOCX).
Structured extraction: spaCy NER + heuristics -> ParsedProfile-shaped JSON
matching the TypeScript interface in server/src/services/ai.ts.

Run: uvicorn main:app --port 8020
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any, Optional

import pdfplumber
import spacy
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from pydantic import BaseModel

from jd_generator import generate_jd

app = FastAPI(title="AIOS Resume Parser", version="1.0.0")

SPACY_MODEL = "en_core_web_sm"
nlp = spacy.load(SPACY_MODEL)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ── Text extraction ──────────────────────────────────────────────────


def extract_pdf_text(data: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=1.5) or ""
            if text.strip():
                parts.append(text)
            # Tables often hold skills/experience grids that extract_text flattens badly
            for table in page.extract_tables():
                for row in table:
                    cells = [c.strip() for c in row if c and c.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


# ── Regex helpers ────────────────────────────────────────────────────

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(
    r"(?:\+\d{1,3}[\s.-]?)?(?:\(\d{2,4}\)[\s.-]?)?\d{3,5}[\s.-]?\d{3,5}(?:[\s.-]?\d{2,5})?"
)
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-./]+", re.I)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-./]+", re.I)
URL_RE = re.compile(r"https?://[^\s|,;]+", re.I)
YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
MONTHS = "Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec"
DATE_TOKEN = rf"(?:(?:{MONTHS})[a-z]*\.?\s+)?(?:19|20)\d{{2}}"
DATE_RANGE_RE = re.compile(
    rf"({DATE_TOKEN})\s*(?:-|–|—|to|until)\s*({DATE_TOKEN}|present|current|now)", re.I
)
EXP_YEARS_RE = re.compile(r"(\d{1,2}(?:\.\d)?)\+?\s*(?:years?|yrs?)", re.I)

SECTION_HEADINGS = {
    "summary": ["summary", "professional summary", "profile", "objective", "about me", "about"],
    "experience": [
        "experience", "work experience", "professional experience", "employment history",
        "employment", "work history", "career history",
    ],
    "education": ["education", "academic background", "academics", "qualifications"],
    "skills": [
        "skills", "technical skills", "core competencies", "technologies",
        "skill set", "key skills", "competencies", "tech stack",
    ],
    "projects": ["projects", "personal projects", "key projects", "academic projects"],
    "certifications": ["certifications", "certificates", "licenses", "courses", "training"],
    "languages": ["languages", "language proficiency"],
}

TECHNICAL_SKILLS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "sql", "nosql", "html",
    "css", "sass", "react", "react native", "angular", "vue", "svelte", "next.js",
    "nuxt", "node.js", "express", "django", "flask", "fastapi", "spring", "spring boot",
    ".net", "rails", "laravel", "graphql", "rest", "grpc", "postgresql", "mysql",
    "mongodb", "redis", "elasticsearch", "cassandra", "dynamodb", "oracle", "sqlite",
    "kafka", "rabbitmq", "spark", "hadoop", "airflow", "aws", "azure", "gcp",
    "google cloud", "docker", "kubernetes", "terraform", "ansible", "jenkins",
    "github actions", "gitlab ci", "ci/cd", "linux", "git", "machine learning",
    "deep learning", "nlp", "computer vision", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "tableau", "power bi", "excel", "jira",
    "figma", "selenium", "cypress", "jest", "pytest", "junit", "microservices",
    "devops", "agile", "scrum", "etl", "data engineering", "data science",
    "android", "ios", "flutter", "salesforce", "sap", "servicenow",
]

SOFT_SKILLS = [
    "leadership", "communication", "teamwork", "problem solving", "problem-solving",
    "time management", "critical thinking", "adaptability", "collaboration",
    "creativity", "attention to detail", "project management", "mentoring",
    "stakeholder management", "negotiation", "public speaking", "decision making",
    "conflict resolution", "analytical",
]

HUMAN_LANGUAGES = [
    "english", "hindi", "spanish", "french", "german", "mandarin", "chinese",
    "japanese", "korean", "arabic", "portuguese", "russian", "italian", "dutch",
    "bengali", "tamil", "telugu", "marathi", "gujarati", "kannada", "malayalam",
    "punjabi", "odia", "urdu",
]

DEGREE_RE = re.compile(
    r"\b(b\.?\s?tech|m\.?\s?tech|b\.?\s?e\b|m\.?\s?e\b|b\.?\s?sc|m\.?\s?sc|b\.?\s?a\b|"
    r"m\.?\s?a\b|b\.?\s?com|m\.?\s?com|bca|mca|mba|pgdm|ph\.?d|bachelor(?:'?s)?|"
    r"master(?:'?s)?|doctorate|diploma|associate)\b[^,\n]*",
    re.I,
)


# ── Section splitting ────────────────────────────────────────────────


def split_sections(text: str) -> dict[str, str]:
    """Split resume text into named sections keyed by canonical heading."""
    heading_lookup: dict[str, str] = {}
    for canonical, variants in SECTION_HEADINGS.items():
        for v in variants:
            heading_lookup[v] = canonical

    lines = text.split("\n")
    sections: dict[str, list[str]] = {"_top": []}
    current = "_top"
    for line in lines:
        stripped = line.strip().strip(":").strip()
        key = stripped.lower()
        # Headings are short standalone lines
        if key in heading_lookup and len(stripped) <= 40:
            current = heading_lookup[key]
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items()}


# ── Field extractors ─────────────────────────────────────────────────


def extract_name(doc: spacy.tokens.Doc, text: str) -> Optional[str]:
    top = text[:300]
    for ent in doc.ents:
        if ent.label_ == "PERSON" and ent.start_char < 300 and 1 <= len(ent.text.split()) <= 4:
            if not EMAIL_RE.search(ent.text) and not any(ch.isdigit() for ch in ent.text):
                return ent.text.strip()
    # Fallback: first short line without contact info
    for line in top.split("\n"):
        line = line.strip()
        if (
            line
            and 1 <= len(line.split()) <= 4
            and not EMAIL_RE.search(line)
            and not any(ch.isdigit() for ch in line)
            and not URL_RE.search(line)
        ):
            return line
    return None


def extract_phone(text: str) -> Optional[str]:
    for m in PHONE_RE.finditer(text[:1500]):
        digits = re.sub(r"\D", "", m.group())
        if 10 <= len(digits) <= 14 and not YEAR_RE.fullmatch(digits):
            return m.group().strip()
    return None


def normalize_url(url: str) -> str:
    return url if url.startswith("http") else f"https://{url}"


def extract_links(text: str) -> dict[str, Optional[str]]:
    linkedin = LINKEDIN_RE.search(text)
    github = GITHUB_RE.search(text)
    portfolio = None
    for m in URL_RE.finditer(text[:2000]):
        u = m.group().rstrip(".,)")
        if "linkedin.com" not in u.lower() and "github.com" not in u.lower():
            portfolio = u
            break
    return {
        "linkedin": normalize_url(linkedin.group().rstrip(".,)")) if linkedin else None,
        "github": normalize_url(github.group().rstrip(".,)")) if github else None,
        "portfolio": portfolio,
    }


SKILL_DISPLAY = {
    "javascript": "JavaScript", "typescript": "TypeScript", "postgresql": "PostgreSQL",
    "mysql": "MySQL", "mongodb": "MongoDB", "nosql": "NoSQL", "graphql": "GraphQL",
    "html": "HTML", "css": "CSS", "sql": "SQL", "aws": "AWS", "gcp": "GCP",
    "php": "PHP", "nlp": "NLP", "etl": "ETL", "sap": "SAP", "rest": "REST",
    "grpc": "gRPC", "ios": "iOS", "devops": "DevOps", "ci/cd": "CI/CD",
    "node.js": "Node.js", "next.js": "Next.js", "react native": "React Native",
    "spring boot": "Spring Boot", "power bi": "Power BI", "github actions": "GitHub Actions",
    "gitlab ci": "GitLab CI", "scikit-learn": "scikit-learn", "pytorch": "PyTorch",
    "tensorflow": "TensorFlow", "dynamodb": "DynamoDB", "rabbitmq": "RabbitMQ",
    "google cloud": "Google Cloud", "junit": "JUnit", "servicenow": "ServiceNow",
}


def skill_display(skill: str) -> str:
    if skill in SKILL_DISPLAY:
        return SKILL_DISPLAY[skill]
    return skill.title() if skill.islower() else skill


def match_skills(text: str, vocabulary: list[str]) -> list[str]:
    lower = text.lower()
    found: list[str] = []
    for skill in vocabulary:
        pattern = r"(?<![\w+#.])" + re.escape(skill) + r"(?![\w+#])"
        if re.search(pattern, lower):
            found.append(skill_display(skill))
    return found


def dedupe_case_insensitive(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            out.append(item)
    return out


def extract_skills_section(section_text: str) -> list[str]:
    items: list[str] = []
    for chunk in re.split(r"[,•|;\n·▪●]+", section_text):
        chunk = re.sub(r"^[\s\-–:*]+|[\s\-–:*]+$", "", chunk)
        # Drop category labels like "Programming Languages:"
        chunk = re.sub(r"^[A-Za-z /&]+:\s*", "", chunk)
        if 1 < len(chunk) <= 40 and not YEAR_RE.search(chunk):
            items.append(chunk)
    return items


def parse_date_token(token: str) -> Optional[datetime]:
    token = token.strip()
    if re.fullmatch(r"(?i)present|current|now", token):
        return datetime.now()
    m = re.search(rf"(?i)({MONTHS})[a-z]*\.?\s+((?:19|20)\d{{2}})", token)
    if m:
        month = [x.lower() for x in MONTHS.split("|")].index(m.group(1).lower()[:3]) + 1
        return datetime(int(m.group(2)), min(month, 12), 1)
    y = YEAR_RE.search(token)
    return datetime(int(y.group()), 1, 1) if y else None


def extract_experience(section_text: str) -> list[dict[str, Any]]:
    if not section_text:
        return []
    entries: list[dict[str, Any]] = []
    lines = [l.strip() for l in section_text.split("\n") if l.strip()]
    for i, line in enumerate(lines):
        m = DATE_RANGE_RE.search(line)
        if not m:
            continue
        header = DATE_RANGE_RE.sub("", line).strip(" -–—|,•:")
        if len(header) < 3 and i > 0:
            header = lines[i - 1].strip(" -–—|,•:")
        title, company = header, ""
        for sep in (" at ", " @ ", " | ", " – ", " — ", " - ", ", "):
            if sep in header:
                title, company = header.split(sep, 1)
                break
        entry = {
            "title": title.strip()[:120] or "Unknown role",
            "company": company.strip()[:120],
            "start_date": m.group(1).strip(),
            "end_date": m.group(2).strip(),
            "description": None,
        }
        # Attach following bullet lines as description
        desc: list[str] = []
        for follow in lines[i + 1 : i + 5]:
            if DATE_RANGE_RE.search(follow):
                break
            desc.append(follow)
        if desc:
            entry["description"] = " ".join(desc)[:500]
        entries.append(entry)
    return entries[:15]


def extract_education(section_text: str, doc: spacy.tokens.Doc) -> list[dict[str, Any]]:
    if not section_text:
        return []
    entries: list[dict[str, Any]] = []
    for line in section_text.split("\n"):
        line = line.strip(" •-–*")
        if not line:
            continue
        deg = DEGREE_RE.search(line)
        if not deg:
            continue
        year = YEAR_RE.findall(line)
        institution = ""
        rest = (line[: deg.start()] + " " + line[deg.end():]).strip(" ,|-–")
        inst_match = re.search(
            r"[A-Z][\w.&' ]*(?:University|Institute|College|School|Academy|IIT|NIT|IIM)[\w.&' ]*",
            line,
        )
        if inst_match:
            institution = inst_match.group().strip()
        elif rest:
            institution = re.sub(r"(19|20)\d{2}", "", rest).strip(" ,|-–")[:100]
        entries.append({
            "degree": deg.group().strip()[:120],
            "institution": institution[:120],
            "year": year[-1] if year else None,
        })
    return entries[:10]


def extract_projects(section_text: str) -> list[dict[str, Any]]:
    if not section_text:
        return []
    projects: list[dict[str, Any]] = []
    lines = [l.strip(" •-–*") for l in section_text.split("\n") if l.strip()]
    current: Optional[dict[str, Any]] = None
    for line in lines:
        # Short-ish line that looks like a title starts a new project
        if len(line) <= 70 and not line.endswith((".", ",")) and len(line.split()) <= 8:
            if current:
                projects.append(current)
            current = {"name": line[:120], "description": None, "technologies": []}
        elif current:
            current["description"] = ((current["description"] or "") + " " + line).strip()[:400]
    if current:
        projects.append(current)
    for p in projects:
        blob = f"{p['name']} {p.get('description') or ''}"
        p["technologies"] = match_skills(blob, TECHNICAL_SKILLS)[:10]
    return projects[:10]


def extract_certifications(section_text: str) -> list[dict[str, Any]]:
    if not section_text:
        return []
    certs: list[dict[str, Any]] = []
    for line in section_text.split("\n"):
        line = line.strip(" •-–*")
        if not (3 <= len(line) <= 150):
            continue
        year = YEAR_RE.search(line)
        name = re.sub(r"[|,(]?\s*(19|20)\d{2}\)?\s*$", "", line).strip(" -–|,")
        certs.append({"name": name[:150], "issuer": None, "date": year.group() if year else None})
    return certs[:15]


def extract_claimed_experience(text: str, summary: str = "") -> Optional[float]:
    """Explicit YOE claim from summary / top-of-resume (skips education phrases)."""
    blobs = [summary or "", (text or "")[:2500]]
    for blob in blobs:
        for m in EXP_YEARS_RE.finditer(blob):
            around = blob[max(0, m.start() - 40) : m.end() + 40].lower()
            if re.search(
                r"bachelor|master|degree|cgpa|gpa|university|college|school|diploma|graduation",
                around,
            ):
                continue
            years = float(m.group(1))
            if 0 <= years <= 50:
                return round(years, 1)
    return None


def employment_years_from_roles(
    experience: list[dict[str, Any]],
) -> tuple[Optional[float], Optional[float]]:
    """Return (calendar_years, sum_years) from employment date spans."""
    intervals: list[tuple[datetime, datetime]] = []
    sum_days = 0.0
    for e in experience:
        start = parse_date_token(e.get("start_date") or "")
        end = parse_date_token(e.get("end_date") or "")
        if start and end and end >= start:
            days = (end - start).days
            sum_days += days
            intervals.append((start, end))
    if not intervals:
        return None, None
    intervals.sort(key=lambda x: x[0])
    merged_days = 0.0
    cur_s, cur_e = intervals[0]
    for s, e in intervals[1:]:
        if s <= cur_e:
            if e > cur_e:
                cur_e = e
        else:
            merged_days += (cur_e - cur_s).days
            cur_s, cur_e = s, e
    merged_days += (cur_e - cur_s).days
    return round(merged_days / 365.25, 1), round(sum_days / 365.25, 1)


def extract_total_experience(
    text: str, experience: list[dict[str, Any]], summary: str = ""
) -> Optional[float]:
    """Prefer employment calendar years; fall back to claimed summary YOE."""
    employment, _sum_years = employment_years_from_roles(experience)
    if employment is not None:
        return employment
    claimed = extract_claimed_experience(text, summary)
    if claimed is not None:
        return claimed
    return None


def experience_consistency(
    text: str,
    experience: list[dict[str, Any]],
    summary: str = "",
    tolerance: float = 1.0,
) -> dict[str, Any]:
    employment, sum_years = employment_years_from_roles(experience)
    claimed = extract_claimed_experience(text, summary)
    mismatch = False
    reason = None
    delta = None
    if claimed is not None and employment is not None:
        delta = round(abs(claimed - employment), 1)
        if delta > tolerance:
            mismatch = True
            reason = (
                f"Experience mismatch: summary claims {claimed} years, "
                f"but employment history totals {employment} years"
            )
    return {
        "employment_years": employment,
        "employment_years_sum": sum_years,
        "claimed_years": claimed,
        "mismatch": mismatch,
        "mismatch_delta": delta,
        "reason": reason,
    }


def compute_confidence(profile: dict[str, Any]) -> float:
    checks = [
        profile.get("name"),
        profile.get("email"),
        profile.get("phone"),
        profile.get("skills"),
        profile.get("experience"),
        profile.get("education"),
        profile.get("professional_summary"),
    ]
    filled = sum(1 for v in checks if (len(v) > 0 if isinstance(v, list) else bool(v)))
    # spaCy heuristics are less certain than LLM parsing; cap below 0.85
    return round(min(0.85, filled / len(checks)), 2)


# ── Profile assembly ─────────────────────────────────────────────────


def build_profile(text: str) -> Optional[dict[str, Any]]:
    doc = nlp(text[:20000])
    sections = split_sections(text)

    name = extract_name(doc, text)
    if not name:
        return None

    links = extract_links(text)
    email = EMAIL_RE.search(text)

    skills_section = sections.get("skills", "")
    section_skills = extract_skills_section(skills_section) if skills_section else []
    technical = dedupe_case_insensitive(match_skills(text, TECHNICAL_SKILLS))
    soft = dedupe_case_insensitive(match_skills(text, SOFT_SKILLS))
    all_skills = dedupe_case_insensitive(section_skills + technical)[:40]

    experience = extract_experience(sections.get("experience", "") or text)
    education = extract_education(sections.get("education", ""), doc)

    current_company = None
    for e in experience:
        if e.get("end_date") and re.fullmatch(r"(?i)present|current|now", e["end_date"]):
            current_company = e.get("company") or None
            break

    summary = sections.get("summary", "")
    languages = [
        lang.title() for lang in HUMAN_LANGUAGES
        if re.search(rf"\b{lang}\b", (sections.get("languages") or text[-1500:]).lower())
    ]

    profile: dict[str, Any] = {
        "name": name,
        "email": email.group() if email else None,
        "phone": extract_phone(text),
        "linkedin": links["linkedin"],
        "github": links["github"],
        "portfolio": links["portfolio"],
        "current_company": current_company,
        "previous_companies": [
            e["company"] for e in experience if e.get("company") and e.get("company") != current_company
        ][:10],
        "experience": experience,
        "education": education,
        "projects": extract_projects(sections.get("projects", "")),
        "skills": all_skills,
        "technical_skills": technical[:30],
        "soft_skills": soft[:15],
        "certifications": extract_certifications(sections.get("certifications", "")),
        "current_salary": None,
        "expected_salary": None,
        "notice_period": None,
        "current_location": None,
        "preferred_location": None,
        "languages": languages[:10],
        "professional_summary": summary[:600] or None,
        "total_experience_years": extract_total_experience(text, experience, summary),
        "experience_consistency": experience_consistency(text, experience, summary),
        "confidence": 0.0,
    }

    # Location: first GPE entity near the top of the resume
    for ent in doc.ents:
        if ent.label_ == "GPE" and ent.start_char < 600:
            profile["current_location"] = ent.text
            break

    notice = re.search(r"notice period[:\s]*([^\n,]{2,40})", text, re.I)
    if notice:
        profile["notice_period"] = notice.group(1).strip()

    profile["confidence"] = compute_confidence(profile)
    return profile


# ── Routes ───────────────────────────────────────────────────────────


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "spacy_model": SPACY_MODEL}


@app.post("/parse")
async def parse(file: UploadFile = File(...)) -> dict[str, Any]:
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    content_type = file.content_type or ""
    filename = (file.filename or "").lower()

    if content_type == PDF_MIME or filename.endswith(".pdf"):
        try:
            text = extract_pdf_text(data)
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"PDF extraction failed: {exc}") from exc
    elif content_type == DOCX_MIME or filename.endswith(".docx"):
        try:
            text = extract_docx_text(data)
        except Exception as exc:
            raise HTTPException(status_code=422, detail=f"DOCX extraction failed: {exc}") from exc
    else:
        # Legacy .doc not supported by python-docx; Node falls back to mammoth
        raise HTTPException(status_code=415, detail="Unsupported file type. Use PDF or DOCX.")

    if not text.strip():
        return {"text": "", "profile": None, "engine": "spacy"}

    return {"text": text, "profile": build_profile(text), "engine": "spacy"}


# ── Job description generation ───────────────────────────────────────


class JDRequest(BaseModel):
    title: str
    client: Optional[str] = None
    location: Optional[str] = None
    open_positions: Optional[int] = None
    notes: Optional[str] = None


@app.post("/generate-jd")
def generate_jd_route(req: JDRequest) -> dict[str, str]:
    if not req.title.strip():
        raise HTTPException(status_code=400, detail="Title required")
    description = generate_jd(
        title=req.title,
        client=req.client,
        location=req.location,
        open_positions=req.open_positions,
        notes=req.notes,
    )
    return {"description": description, "engine": "template"}
